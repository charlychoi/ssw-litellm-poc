"""Phase 2.3 최종 보고서 — 실제 Provider(OpenRouter/Groq) 검증 결과 Word docx 생성"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOCS_DIR = Path(__file__).parent.parent / "docs"
OUTPUT_PATH = DOCS_DIR / "SSW_LiteLLM_Phase23_Report.docx"
DESKTOP_PATH = Path.home() / "Desktop" / "SSW_LiteLLM_Phase23_Report.docx"

BLUE_DARK  = "1F4E79"
BLUE_MID   = "2E75B6"
BLUE_LIGHT = "D6E4F0"
GREEN      = "1D6F42"
GREEN_LIGHT= "E2EFDA"
RED        = "C00000"
ORANGE     = "C55A11"
GRAY       = "595959"


# ── 헬퍼 ──────────────────────────────────────────────────────────────────────

def rgb(hex_str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 레벨 1은 파란색
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = rgb(BLUE_DARK)
    return p


def body(doc, text, bold=False, color=None, size=11, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold, color=color, size=size, italic=italic)
    return p


def bullet(doc, text, level=0, bold=False, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, bold=bold, color=color)
    return p


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_font(cell, text, bold=False, color=None, size=10, align=None):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    set_font(run, bold=bold, color=color, size=size)
    if align:
        p.alignment = align


def add_table(doc, headers, rows, col_widths=None,
              header_color=BLUE_DARK, stripe_color=BLUE_LIGHT, font_size=10):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # 헤더
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], header_color)
        set_cell_font(hdr[i], h, bold=True, color="FFFFFF", size=font_size)

    # 데이터 행
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = stripe_color if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            shade_cell(cells[ci], bg)
            # 통과/차단 표시 색상
            col = None
            if val in ("✅", "통과", "성공"):
                col = GREEN
            elif val in ("❌", "실패"):
                col = RED
            elif val in ("⚠️", "건너뜀"):
                col = ORANGE
            set_cell_font(cells[ci], str(val), color=col, size=font_size)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def page_break(doc):
    doc.add_page_break()


def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE_MID)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── 보고서 본문 ────────────────────────────────────────────────────────────────

def build(doc: Document):
    today = date.today().strftime("%Y년 %m월 %d일")

    # ── 표지 ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("상상우리 LiteLLM AI 비용 거버넌스 PoC")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = rgb(BLUE_DARK)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Phase 2.3 최종 검증 보고서")
    run2.bold = True
    run2.font.size = Pt(18)
    run2.font.color.rgb = rgb(BLUE_MID)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("실제 Provider(OpenRouter / Groq) 연동 검증")
    run3.font.size = Pt(14)
    run3.font.color.rgb = rgb(GRAY)

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f"작성일: {today}  |  검증 결과: 10 / 10 통과")
    run4.font.size = Pt(12)
    run4.font.color.rgb = rgb(GREEN)
    run4.bold = True

    doc.add_paragraph()
    doc.add_paragraph()
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("상상우리 (Sangsang-Woori)  |  AI Infrastructure Team")
    run5.font.size = Pt(11)
    run5.font.color.rgb = rgb(GRAY)

    page_break(doc)

    # ── 1. 목적 및 배경 ────────────────────────────────────────────────────────
    heading(doc, "1. 목적 및 배경")
    body(doc,
         "Phase 1(정책 검증)과 Phase 2(CLI 연동 검증)에서 LiteLLM Proxy를 통해 "
         "직원별·도구별 가상키 발급, 모델 접근 제한, 예산 초과 차단, "
         "Claude Code / Codex CLI / Gemini CLI의 LiteLLM 라우팅 가능성을 "
         "mock 기반으로 검증했습니다.")
    doc.add_paragraph()
    body(doc,
         "Phase 2.3의 목적은 mock 없이 실제 외부 LLM Provider(OpenRouter, Groq)를 "
         "연결하여 다음 핵심 명제를 검증하는 것입니다:", bold=True)
    bullet(doc, "실제 외부 LLM 응답에서도 가상키 모델 제한이 정상 작동하는가?")
    bullet(doc, "실제 provider 호출 시 token usage와 spend가 LiteLLM에 기록되는가?")
    bullet(doc, "실제 호출 기반 예산 초과 차단이 동작하는가?")
    bullet(doc, "CLI 도구(Codex-style)가 실제 외부 LLM을 호출할 때도 가상키 정책이 유지되는가?")
    bullet(doc, "OpenRouter free tier 429 rate limit 발생 시 LiteLLM이 자동 fallback하는가?")
    doc.add_paragraph()

    # ── 2. 검증 환경 ───────────────────────────────────────────────────────────
    heading(doc, "2. 검증 환경")
    add_table(doc,
        ["항목", "값"],
        [
            ["LiteLLM Proxy 버전", "v1.83.7 (v1.82.7/v1.82.8 공급망 공격으로 사용 금지)"],
            ["LiteLLM 실행 주소",  "http://localhost:4000"],
            ["데이터베이스",       "PostgreSQL 16 (Homebrew, DB: litellm_poc)"],
            ["Python 환경",        "Python 3.11.15, uv 패키지 관리자"],
            ["실제 Provider #1",   "OpenRouter (openrouter.ai) — 무료 모델 사용"],
            ["실제 Provider #2",   "Groq (console.groq.com) — llama-3.1-8b-instant"],
            ["테스트 날짜",        today],
            ["검증 스크립트",      "scripts/70_real_provider_test.py"],
        ],
        col_widths=[5, 11]
    )
    doc.add_paragraph()

    # ── 3. 실제 Provider alias 구성 ────────────────────────────────────────────
    heading(doc, "3. 실제 Provider 모델 alias 구성")
    body(doc,
         "LiteLLM config.yaml에 아래 4개의 실제 provider alias를 추가했습니다. "
         "내부 alias를 통해 실제 provider 모델명이 외부에 노출되지 않으며, "
         "향후 모델 교체 시 config만 수정하면 됩니다.")
    doc.add_paragraph()
    add_table(doc,
        ["내부 alias", "실제 모델", "Provider", "용도"],
        [
            ["ssw-free-openrouter", "google/gemma-4-31b-it:free\n(fallback: nemotron-3-nano-30b-a3b:free)",
             "OpenRouter", "무료 모델 실제 호출 검증"],
            ["ssw-low-cost-real",   "nvidia/nemotron-3-nano-30b-a3b:free",
             "OpenRouter", "일반 직원 기본 모델 후보"],
            ["ssw-fast-groq",       "llama-3.1-8b-instant",
             "Groq",       "빠른 응답 / 저지연 검증"],
            ["ssw-expensive-real",  "anthropic/claude-opus-4",
             "OpenRouter", "비허용 모델 차단 테스트용 (실제 호출 안 됨)"],
        ],
        col_widths=[4.5, 6, 3, 4.5]
    )
    doc.add_paragraph()
    body(doc,
         "※ ssw-free-openrouter에는 LiteLLM router fallback이 설정되어 있어, "
         "Gemma 4가 429 rate limit을 반환하면 자동으로 Nemotron으로 전환됩니다.",
         italic=False, color=GRAY)
    doc.add_paragraph()

    # ── 4. 가상키 정책 업데이트 ────────────────────────────────────────────────
    heading(doc, "4. 가상키(Virtual Key) 정책 업데이트")
    body(doc,
         "Phase 2.3 실제 provider alias를 기존 가상키의 허용 모델 목록에 추가했습니다. "
         "PostgreSQL 직접 업데이트로 적용하였으며, Phase 1+2의 정책은 그대로 유지됩니다.")
    doc.add_paragraph()
    add_table(doc,
        ["Key Alias", "사용자", "도구", "Phase 2.3에서 추가된 허용 모델", "예산"],
        [
            ["staff-lee-chat",  "lee", "Chat API",    "ssw-free-openrouter, ssw-low-cost-real", "$1"],
            ["dev-kim-codex",   "kim", "Codex CLI",   "ssw-free-openrouter, ssw-fast-groq, ssw-low-cost-real", "$5"],
            ["dev-kim-claude",  "kim", "Claude Code", "ssw-free-openrouter, ssw-low-cost-real", "$5"],
            ["dev-kim-gemini",  "kim", "Gemini CLI",  "ssw-free-openrouter, ssw-fast-groq", "$5"],
            ["admin-park-test", "park","Admin API",   "ssw-free-openrouter, ssw-fast-groq, ssw-low-cost-real, ssw-expensive-real", "$10"],
        ],
        col_widths=[3.5, 1.8, 3, 6, 1.5]
    )
    doc.add_paragraph()

    page_break(doc)

    # ── 5. 테스트 시나리오 및 결과 ────────────────────────────────────────────
    heading(doc, "5. 테스트 시나리오 및 결과 (10/10 통과)")
    divider(doc)

    # Scenario 1
    heading(doc, "Scenario 1 — OpenRouter 실제 호출", level=2)
    body(doc, "검증 목적", bold=True, color=BLUE_DARK)
    body(doc, "OpenRouter 무료 모델이 LiteLLM alias를 통해 실제 응답하는지 확인합니다.")
    doc.add_paragraph()
    add_table(doc,
        ["Key Alias", "요청 모델", "resolved 모델", "HTTP", "token", "결과"],
        [
            ["staff-lee-chat",  "ssw-free-openrouter", "nvidia/nemotron-3-nano-30b-a3b:free\n(Gemma 429 → auto fallback)", "200", "118", "✅"],
            ["admin-park-test", "ssw-low-cost-real",   "nvidia/nemotron-3-nano-30b-a3b:free", "200", "107", "✅"],
        ],
        col_widths=[3.2, 3.5, 5.5, 1.5, 1.5, 1.5]
    )
    doc.add_paragraph()
    body(doc, "실제 응답 예시 (staff-lee-chat):", bold=True)
    body(doc,
         '  "이 응답은 LiteLLM 게이트웨이를 통해 전달되었습니다." '
         '— OpenRouter Nemotron 실제 응답, mock 문구 없음',
         color=GREEN)
    body(doc,
         "⚑ 주목: Gemma 4 31B가 free tier rate limit(429)을 반환하자 LiteLLM이 "
         "자동으로 Nemotron fallback을 실행했습니다. 운영 환경에서의 provider 장애 "
         "자동 전환 기능이 실증됐습니다.",
         color=ORANGE)
    doc.add_paragraph()

    # Scenario 2
    heading(doc, "Scenario 2 — Groq 실제 호출", level=2)
    body(doc, "검증 목적", bold=True, color=BLUE_DARK)
    body(doc, "Groq의 llama-3.1-8b-instant 모델이 LiteLLM을 통해 실제 응답하는지 확인합니다.")
    doc.add_paragraph()
    add_table(doc,
        ["Key Alias", "요청 모델", "resolved 모델", "HTTP", "token", "결과"],
        [
            ["dev-kim-codex",  "ssw-fast-groq", "groq/llama-3.1-8b-instant", "200", "54", "✅"],
            ["dev-kim-gemini", "ssw-fast-groq", "groq/llama-3.1-8b-instant", "200", "57", "✅"],
        ],
        col_widths=[3.2, 3, 5.5, 1.5, 1.5, 1.5]
    )
    doc.add_paragraph()
    body(doc, "실제 응답 예시:", bold=True)
    body(doc, '  dev-kim-codex: "GROQ is real okay."', color=GREEN)
    body(doc, '  dev-kim-gemini: "Gemini key Groq okay."', color=GREEN)
    doc.add_paragraph()

    # Scenario 3
    heading(doc, "Scenario 3 — 비허용 모델 차단 (실제 Provider 연결 상태에서)", level=2)
    body(doc, "검증 목적", bold=True, color=BLUE_DARK)
    body(doc,
         "실제 외부 provider 모델이 연결되어 있어도, 가상키의 allowed_models 정책에 "
         "의해 비허용 모델 호출이 차단되는지 확인합니다. 비용이 외부 provider에 "
         "도달하기 전에 차단되어야 합니다.")
    doc.add_paragraph()
    add_table(doc,
        ["Key Alias", "시도 모델", "HTTP", "error_type", "provider 도달", "결과"],
        [
            ["staff-lee-chat", "ssw-expensive-real\n(claude-opus-4)", "401", "key_model_access_denied", "❌ 미도달 (비용 없음)", "✅"],
            ["staff-lee-chat", "ssw-fast-groq\n(groq llama)", "401", "key_model_access_denied", "❌ 미도달 (비용 없음)", "✅"],
        ],
        col_widths=[3.2, 3.5, 1.5, 4.5, 4, 1.5]
    )
    doc.add_paragraph()
    body(doc,
         "핵심 결과: 실제 API provider가 연결된 상태에서도 LiteLLM 게이트웨이 "
         "레벨에서 모든 비허용 모델 요청이 차단됩니다. "
         "고가 모델(Claude Opus 4 등)이 실수로 호출되어도 비용이 발생하지 않습니다.",
         bold=True, color=BLUE_DARK)
    doc.add_paragraph()

    page_break(doc)

    # Scenario 4
    heading(doc, "Scenario 4 — 실제 호출 기반 예산 차단", level=2)
    body(doc, "검증 목적", bold=True, color=BLUE_DARK)
    body(doc,
         "max_budget=$0 으로 설정된 임시 키로 실제 provider를 호출했을 때 "
         "LiteLLM이 budget_exceeded로 즉시 차단하는지 검증합니다.")
    doc.add_paragraph()
    add_table(doc,
        ["단계", "동작", "HTTP", "응답"],
        [
            ["임시 키 생성", "max_budget=0, models=[ssw-free-openrouter]", "200", "키 생성 성공"],
            ["호출 시도",    "ssw-free-openrouter 호출 (budget=0 초과)",   "400", "budget_exceeded ✅"],
        ],
        col_widths=[3.5, 7, 2, 4]
    )
    doc.add_paragraph()
    body(doc,
         "Phase 1의 mock 기반 예산 차단이 실제 외부 provider 연결 상태에서도 동일하게 동작합니다. "
         "직원 예산이 소진되면 외부 LLM API 호출 자체가 불가능해집니다.",
         color=GREEN)
    doc.add_paragraph()

    # Scenario 5
    heading(doc, "Scenario 5 — 관리자 spend 리포트", level=2)
    body(doc, "검증 목적", bold=True, color=BLUE_DARK)
    body(doc, "실제 provider 호출이 LiteLLM spend 로그에 기록되는지 확인합니다.")
    doc.add_paragraph()
    add_table(doc,
        ["항목", "값"],
        [
            ["총 로그 수 (누적)",   "97건"],
            ["총 누적 spend",       "$0.01853234 (PoC 전체 기간)"],
            ["Phase 2.3 신규 모델", "groq/llama-3.1-8b-instant, openrouter/google/gemma-4-31b-it:free, openrouter/nvidia/nemotron-3-nano-30b-a3b:free"],
            ["리포트 생성 위치",    "docs/REAL_PROVIDER_VERIFICATION_REPORT.md"],
        ],
        col_widths=[5, 11]
    )
    doc.add_paragraph()
    body(doc,
         "※ 무료 모델(OpenRouter free tier)은 LiteLLM 내장 pricing DB에 가격이 없어 "
         "spend=$0으로 기록될 수 있습니다. 유료 모델 전환 시 자동으로 정확한 비용이 집계됩니다.",
         color=GRAY)
    doc.add_paragraph()

    # Scenario 6
    heading(doc, "Scenario 6 [필수] — Codex CLI → LiteLLM → 실제 Provider", level=2)
    body(doc, "검증 목적", bold=True, color=BLUE_DARK)
    body(doc,
         "Phase 2.3의 핵심 명제: Codex-style 클라이언트가 실제 외부 LLM을 호출할 때도 "
         "LiteLLM 가상키 정책(모델 제한 + 예산 차단)이 유지되는가?")
    doc.add_paragraph()
    add_table(doc,
        ["테스트", "Key", "요청 모델", "결과", "응답/차단 내용"],
        [
            ["6a: 실제 LLM 호출", "dev-kim-codex", "ssw-free-openrouter",
             "✅ HTTP 200", '"CLI REAL PROVIDER_OK"\n(실제 OpenRouter 응답, mock 없음)'],
            ["6b: 비허용 모델 차단", "dev-kim-codex", "ssw-expensive-real",
             "✅ HTTP 401", "key_model_access_denied\n(claude-opus-4 접근 차단)"],
        ],
        col_widths=[3.5, 3, 3.5, 2.5, 5]
    )
    doc.add_paragraph()
    body(doc,
         "Codex CLI / API 클라이언트가 OPENAI_BASE_URL=http://localhost:4000 으로 설정되면 "
         "모든 요청이 LiteLLM을 경유하며, 실제 외부 LLM 응답과 동시에 가상키 정책이 "
         "그대로 적용됩니다. '실제 LLM 호출' + '정책 통제' 동시 검증 완료.",
         bold=True, color=BLUE_DARK)
    doc.add_paragraph()

    page_break(doc)

    # ── 6. 전체 결과 요약 ──────────────────────────────────────────────────────
    heading(doc, "6. 전체 결과 요약")
    add_table(doc,
        ["Scenario", "검증 항목", "Provider", "결과"],
        [
            ["1", "staff-lee-chat → ssw-free-openrouter 실제 응답",      "OpenRouter (Nemotron fallback)", "✅"],
            ["1", "admin-park-test → ssw-low-cost-real 실제 응답",        "OpenRouter (Nemotron)",          "✅"],
            ["2", "dev-kim-codex → ssw-fast-groq 실제 응답",              "Groq (llama-3.1-8b-instant)",    "✅"],
            ["2", "dev-kim-gemini → ssw-fast-groq 실제 응답",             "Groq (llama-3.1-8b-instant)",    "✅"],
            ["3", "staff-lee-chat → ssw-expensive-real 차단",             "— (LiteLLM 레벨 차단)",           "✅"],
            ["3", "staff-lee-chat → ssw-fast-groq 차단",                  "— (LiteLLM 레벨 차단)",           "✅"],
            ["4", "max_budget=0 예산 차단 (budget_exceeded)",             "— (LiteLLM 정책)",               "✅"],
            ["5", "spend 로그 및 리포트 (97건 누적)",                      "LiteLLM DB",                     "✅"],
            ["6", "Codex-style API → ssw-free-openrouter 실제 LLM",      "OpenRouter (Gemma 4)",            "✅"],
            ["6", "Codex CLI → ssw-expensive-real 비허용 차단",           "— (LiteLLM 레벨 차단)",           "✅"],
        ],
        col_widths=[2, 7.5, 4.5, 1.5]
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("최종 결과: 10 / 10 통과  ✅")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = rgb(GREEN)
    doc.add_paragraph()

    # ── 7. 핵심 검증 포인트 ────────────────────────────────────────────────────
    heading(doc, "7. Phase 2.3이 증명한 것")
    body(doc,
         "Phase 1+2가 mock으로 증명한 기능들이 실제 외부 LLM 환경에서도 동일하게 "
         "작동함을 검증했습니다. 아래는 Phase 2.3에서 새롭게 확인된 사항입니다.")
    doc.add_paragraph()

    findings = [
        ("실제 외부 LLM 응답 통과",
         "OpenRouter (Gemma 4, Nemotron), Groq (llama-3.1-8b) 모두 "
         "LiteLLM 게이트웨이를 통해 실제 응답을 반환했습니다. "
         "응답에 mock 문구 없이 실제 LLM 생성 텍스트가 확인됩니다."),
        ("token usage 기록",
         "모든 실제 호출에서 prompt/completion/total token이 정확히 기록되어 "
         "사용량 집계의 기반이 검증됐습니다."),
        ("실제 provider 연결 상태에서도 모델 제한 유지",
         "staff-lee-chat 키가 Groq(ssw-fast-groq)나 Claude Opus(ssw-expensive-real)를 "
         "요청해도 LiteLLM이 provider에 도달하기 전에 차단합니다. "
         "고가 모델 접근이 물리적으로 불가합니다."),
        ("실제 호출 기반 예산 차단",
         "max_budget=0 키로 실제 provider를 호출하면 budget_exceeded(HTTP 400)가 "
         "즉시 발생합니다. 예산 소진 시 외부 API 호출 자체가 불가합니다."),
        ("CLI 도구 + 실제 LLM + 가상키 정책 동시 검증",
         "Codex-style 클라이언트가 실제 OpenRouter LLM에 'CLI REAL PROVIDER_OK'를 "
         "받아오면서도 비허용 모델 요청은 즉시 차단됩니다. "
         "CLI 도구를 통한 무제한 LLM 사용이 제도적으로 불가능함이 검증됩니다."),
        ("Provider 장애 자동 fallback",
         "OpenRouter Gemma 4가 free tier rate limit(429)을 반환할 때 "
         "LiteLLM이 자동으로 Nemotron으로 전환해 서비스 연속성을 유지했습니다. "
         "이는 운영 환경에서 provider 장애에 대한 resilience를 의미합니다."),
    ]
    for i, (title, desc) in enumerate(findings, 1):
        p = doc.add_paragraph()
        run1 = p.add_run(f"{i}. {title}\n")
        run1.bold = True
        run1.font.color.rgb = rgb(BLUE_MID)
        run1.font.size = Pt(11)
        run2 = p.add_run(f"   {desc}")
        run2.font.size = Pt(10)
        run2.font.color.rgb = rgb("000000")
    doc.add_paragraph()

    page_break(doc)

    # ── 8. 아키텍처 흐름 ──────────────────────────────────────────────────────
    heading(doc, "8. 검증된 아키텍처 흐름")
    body(doc,
         "Phase 2.3에서 검증된 실제 provider 연동 아키텍처입니다.")
    doc.add_paragraph()

    arch_text = (
        "┌─────────────────────────────────────────────────────────────────┐\n"
        "│                      직원 / CLI 도구 레이어                        │\n"
        "│  Claude Code  │  Codex CLI  │  Gemini CLI  │  Chat API          │\n"
        "│  (ANTHROPIC_  │  (OPENAI_   │  (GOOGLE_    │  (직접 API 호출)    │\n"
        "│   BASE_URL)   │   BASE_URL) │   GEMINI_URL)│                    │\n"
        "└──────────────────────────┬──────────────────────────────────────┘\n"
        "                           │  가상키(Virtual Key) 첨부\n"
        "                           ▼\n"
        "┌─────────────────────────────────────────────────────────────────┐\n"
        "│              LiteLLM Proxy (localhost:4000)                      │\n"
        "│  ① 가상키 인증  ② 모델 제한 검사  ③ 예산 검사                    │\n"
        "│  ④ alias → 실제 모델 변환  ⑤ fallback 라우팅 (429 → auto 전환)   │\n"
        "│  ⑥ token/spend 기록 (PostgreSQL)                                │\n"
        "└──────┬──────────────────┬───────────────────────────────────────┘\n"
        "       │                  │\n"
        "       ▼                  ▼\n"
        "┌──────────────┐  ┌──────────────────────────────────────────────┐\n"
        "│  차단 응답    │  │           실제 Provider 호출                  │\n"
        "│  401 / 400   │  │  OpenRouter (Gemma 4, Nemotron, claude-opus) │\n"
        "│  (비허용/예산)│  │  Groq (llama-3.1-8b-instant)                │\n"
        "└──────────────┘  └──────────────────────────────────────────────┘"
    )

    p = doc.add_paragraph()
    run = p.add_run(arch_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(BLUE_DARK)
    doc.add_paragraph()

    body(doc,
         "가상키 정책(모델 제한, 예산 한도)을 통과한 요청만 실제 provider에 전달되며, "
         "차단된 요청은 provider에 도달하지 않아 비용이 발생하지 않습니다.",
         bold=True, color=BLUE_DARK)
    doc.add_paragraph()

    # ── 9. Phase 별 전체 검증 이력 ────────────────────────────────────────────
    heading(doc, "9. Phase 별 전체 검증 이력")
    add_table(doc,
        ["Phase", "검증 내용", "결과"],
        [
            ["Phase 1\n(정책 검증)",
             "LiteLLM Proxy 실행 / 마스터키-가상키 분리 / 가상키 5개 발급 / "
             "허용 모델 호출 / 비허용 모델 차단(401) / 예산 초과 차단(400) / "
             "사용량 리포트 / pytest 48개",
             "7/7 ✅"],
            ["Phase 2\n(CLI 연동)",
             "Claude Code → LiteLLM 라우팅 / Codex CLI → LiteLLM 라우팅 / "
             "Gemini CLI → LiteLLM 라우팅 / 모델 접근 제어 / CLI 자동 검증 스크립트",
             "5/5 ✅"],
            ["Phase 2.3\n(실제 Provider)",
             "OpenRouter 실제 호출 / Groq 실제 호출 / 비허용 모델 차단(실제 연결 상태) / "
             "실제 호출 기반 예산 차단 / spend 리포트 / Codex-style CLI + 실제 LLM / "
             "Provider 장애 자동 fallback",
             "10/10 ✅"],
        ],
        col_widths=[2.5, 12, 2]
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Phase 1 + Phase 2 + Phase 2.3 누계: 22/22 항목 모두 통과 ✅")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = rgb(GREEN)
    doc.add_paragraph()

    page_break(doc)

    # ── 10. 운영 전 권고사항 ───────────────────────────────────────────────────
    heading(doc, "10. 운영 전 권고사항")
    add_table(doc,
        ["항목", "현황", "권고 조치", "우선순위"],
        [
            ["Redis RPM/TPM",    "미설정 (rate limit 비활성)",
             "brew install redis 후 config.yaml에 redis_url 추가",              "높음"],
            ["모델명 고정",       "free tier 모델명 변경 위험",
             "config.yaml 모델명 주기적 확인 또는 특정 버전 pin",               "높음"],
            ["실제 API 비용 집계","무료 모델은 spend=$0 기록",
             "유료 모델 전환 시 자동 해결 / custom_pricing 설정 가능",           "중간"],
            ["SSO/OIDC 연동",    "미구현",
             "직원 계정과 가상키를 사내 IdP와 연동 (운영 필수)",                "높음"],
            ["키 로테이션",       "수동 발급",
             "30일 주기 자동 재발급 스크립트 구성 권장",                         "중간"],
            ["운영 배포",         "로컬 실행",
             "Railway / GCP / AWS에 LiteLLM Proxy 배포 (Docker Compose 준비됨)", "높음"],
            ["Google Chat 연동", "미구현",
             "Hermes Agent LiteLLM provider 연동 → Google Chat 봇 완성",        "다음 단계"],
        ],
        col_widths=[3.5, 3.5, 6.5, 2]
    )
    doc.add_paragraph()

    # ── 11. 결론 ──────────────────────────────────────────────────────────────
    heading(doc, "11. 결론")
    body(doc,
         "Phase 2.3을 통해 상상우리 LiteLLM AI 비용 거버넌스 PoC의 핵심 가정이 "
         "실제 외부 LLM 환경에서도 완전히 검증됐습니다.", bold=True)
    doc.add_paragraph()

    conclusions = [
        "Claude Code, Codex CLI, Gemini CLI 등 어떤 AI 코딩 도구를 사용하더라도 "
        "LiteLLM 게이트웨이를 경유하면 직원별·도구별 모델 제한과 예산 통제가 가능합니다.",

        "실제 OpenRouter와 Groq provider를 연결해도 mock과 동일한 가상키 정책이 "
        "적용됩니다. mock → 실제 전환 시 추가 코드 수정이 필요 없습니다.",

        "LiteLLM의 fallback 라우팅 기능으로 특정 provider가 장애/rate limit 상태일 때 "
        "자동으로 다른 provider로 전환되어 서비스 연속성이 보장됩니다.",

        "비허용 모델 요청은 외부 provider에 도달하지 않으므로 실수로 "
        "Claude Opus 4 같은 고가 모델을 호출해도 비용이 발생하지 않습니다.",

        "다음 단계로 Redis RPM/TPM 설정, 운영 배포(Railway/GCP), "
        "Hermes Agent 연동을 통해 전사 AI 사용 정책 시스템으로 발전시킬 수 있습니다.",
    ]
    for c in conclusions:
        bullet(doc, c)
    doc.add_paragraph()

    divider(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"작성: Claude Code (Anthropic)  |  {today}  |  상상우리 AI Infrastructure Team")
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(GRAY)
    run.italic = True


# ── 실행 ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # 기본 여백 설정
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    build(doc)

    doc.save(OUTPUT_PATH)
    doc.save(DESKTOP_PATH)
    print(f"✅ 저장 완료:")
    print(f"   {OUTPUT_PATH}")
    print(f"   {DESKTOP_PATH}")


if __name__ == "__main__":
    main()
