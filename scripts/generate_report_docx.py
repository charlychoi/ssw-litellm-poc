"""PoC 결과 보고서 Word(.docx) 생성 스크립트"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOCS_DIR = Path(__file__).parent.parent / "docs"
OUTPUT_PATH = DOCS_DIR / "SSW_LiteLLM_Governance_PoC_Report.docx"


# ── 헬퍼 ──────────────────────────────────────────────────────────────────────

def set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def body(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold, color=color)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # 헤더 행
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)

    # 데이터 행
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
        if ri % 2 == 1:
            for ci in range(len(headers)):
                tc = cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D6E4F0")
                tcPr.append(shd)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


# ── 본문 ──────────────────────────────────────────────────────────────────────

def build_report():
    doc = Document()

    # 기본 스타일 여백
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ── 표지 ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("상상우리 LiteLLM AI 비용 거버넌스")
    set_font(run, bold=True, size=22, color=(31, 78, 121))

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = t2.add_run("Proof of Concept 검증 결과 보고서")
    set_font(run2, bold=True, size=18, color=(31, 78, 121))

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"작성일: {date.today().strftime('%Y년 %m월 %d일')}     "
                 f"버전: 1.0     LiteLLM 1.83.7")

    doc.add_page_break()

    # ── 1. 개요 ───────────────────────────────────────────────────────────────
    heading(doc, "1. 개요")

    body(doc, "1.1 PoC 목적", bold=True)
    body(doc,
         "상상우리 AI 에이전트 프로젝트에서는 직원들이 Claude Code, Codex CLI, Gemini CLI 등 "
         "다양한 유료 LLM 도구를 사용할 예정이다. 이 도구들은 단순 채팅보다 토큰 소비가 많고 "
         "장시간 자동 실행되므로 기업 예산이 통제 없이 초과될 위험이 있다. "
         "본 PoC는 Google Chat 연동 전 단계에서 LiteLLM Proxy를 AI 게이트웨이로 활용하여 "
         "직원별·도구별·모델별 LLM 사용을 예산 범위 안에서 실제로 차단하고 추적할 수 있는지를 검증하였다.")

    body(doc, "1.2 검증 핵심 질문", bold=True)
    for q in [
        "직원 및 도구마다 별도의 가상키(Virtual Key)를 발급하고 식별할 수 있는가?",
        "가상키별로 허용 모델과 금지 모델을 분리할 수 있는가?",
        "월 예산을 설정하고 초과 시 호출을 실제로 차단할 수 있는가?",
        "누가, 어떤 도구로, 어떤 모델을, 얼마나 사용했는지 리포트할 수 있는가?",
        "Claude Code / Codex CLI 같은 코딩 에이전트도 이 게이트웨이로 통제 가능한가?",
    ]:
        bullet(doc, q)

    body(doc, "1.3 기술 스택", bold=True)
    add_table(doc,
        ["영역", "선택", "비고"],
        [
            ["AI Gateway", "LiteLLM Proxy >=1.83.10", "v1.82.7/v1.82.8 공급망 공격 이슈로 사용 금지; >=1.83.10 권장"],
            ["데이터베이스", "PostgreSQL 16", "가상키 및 사용량 로그 저장"],
            ["속도 제한", "Redis (선택)", "미사용 시 RPM/TPM 제한 비활성 — 예산/모델 제한은 정상 작동"],
            ["관리자 UI", "Streamlit", "사용자 테스트 탭 + 관리자 통제 탭"],
            ["테스트", "pytest 48개", "정책 로직 단위 테스트"],
            ["패키지 관리", "uv (Python 3.11)", "의존성 고정 및 재현성 보장"],
        ],
        col_widths=[3, 4, 8],
    )

    doc.add_paragraph()

    # ── 2. 검증 결과 ──────────────────────────────────────────────────────────
    heading(doc, "2. 검증 결과 (7/7 성공)")

    add_table(doc,
        ["#", "검증 항목", "결과", "근거"],
        [
            ["1", "LiteLLM Proxy 실행", "✅ 성공", "curl /health/liveliness → \"I'm alive!\""],
            ["2", "마스터키 / 가상키 분리", "✅ 성공", "직원은 가상키만 보유, 마스터키 미노출"],
            ["3", "직원/도구별 가상키 발급", "✅ 성공", "5개 키 생성, alias/team/모델/예산 개별 설정"],
            ["4", "허용 모델 호출 성공", "✅ 성공", "dev-kim-claude 키 + ssw-fake 모델 → 응답 수신"],
            ["5", "비허용 모델 차단", "✅ 성공", "staff-lee-chat 키 + ssw-expensive → 401 key_model_access_denied"],
            ["6", "예산 초과 차단", "✅ 성공", "max_budget 초과 후 → 400 Budget has been exceeded!"],
            ["7", "사용량 리포트 생성", "✅ 성공", "사용자/도구/모델별 비용 집계 및 Markdown 출력"],
        ],
        col_widths=[1, 4.5, 2, 7.5],
    )

    doc.add_paragraph()

    # ── 3. 세부 검증 내용 ──────────────────────────────────────────────────────
    heading(doc, "3. 세부 검증 내용")

    # 3.1
    heading(doc, "3.1 가상키 발급 및 정책 설정", level=2)
    body(doc,
         "LiteLLM Admin API를 통해 직원 및 도구별로 가상키를 자동 발급하였다. "
         "각 키는 허용 모델 목록, 월 예산, 분당 요청 제한(RPM), 분당 토큰 제한(TPM), "
         "사용자 ID, 팀 ID, 도구 메타데이터를 독립적으로 보유한다.")

    add_table(doc,
        ["키 Alias", "직원", "팀", "도구", "허용 모델", "월 예산"],
        [
            ["dev-kim-claude", "kim", "dev-team", "Claude Code", "ssw-dev-sonnet, ssw-fake", "$5.00"],
            ["dev-kim-codex", "kim", "dev-team", "Codex CLI", "ssw-dev-gpt, ssw-fake", "$5.00"],
            ["dev-kim-gemini", "kim", "dev-team", "Gemini CLI", "ssw-free-test, ssw-fake", "$5.00"],
            ["staff-lee-chat", "lee", "staff-team", "일반 채팅", "ssw-low-cost, ssw-fake", "$1.00"],
            ["admin-park-test", "park", "admin-team", "관리자 API", "전 모델", "$10.00"],
        ],
        col_widths=[3.5, 1.5, 2.5, 2.5, 3.5, 1.5],
    )
    body(doc, "\n중요 보안 원칙:")
    bullet(doc, "마스터키(LITELLM_MASTER_KEY)는 관리자 서버에서만 사용하며 직원에게 절대 공유하지 않는다.")
    bullet(doc, "실제 키 원문은 문서에 저장하지 않으며, 앞 8자만 redacted 형태로 보관한다.")
    bullet(doc, "키 회수, 예산 조정, 모델 권한 변경은 Admin API를 통해 즉시 적용된다.")

    doc.add_paragraph()

    # 3.2
    heading(doc, "3.2 모델 접근 제한 검증", level=2)
    body(doc,
         "일반 직원(lee)의 키 staff-lee-chat은 저가 모델만 허용하도록 설정하였다. "
         "해당 키로 고가 모델(ssw-expensive)을 호출하면 LiteLLM이 즉시 차단한다.")

    body(doc, "실행 명령:", bold=True)
    add_code_block(doc, "uv run python scripts/30_demo_calls.py --user lee --tool chat --case denied-model")

    body(doc, "LiteLLM 응답 (HTTP 401):", bold=True)
    add_code_block(doc,
        '{\n'
        '  "error": {\n'
        '    "message": "key not allowed to access model. This key can only access\n'
        '                models=[\'ssw-low-cost\', \'ssw-fake\', \'low-cost\'].\n'
        '                Tried to access ssw-expensive",\n'
        '    "type": "key_model_access_denied",\n'
        '    "code": "401"\n'
        '  }\n'
        '}'
    )

    body(doc,
         "이 차단은 네트워크 레벨에서 발생하며, 잘못된 모델로의 호출이 provider(Anthropic/OpenAI)에 "
         "도달하기 전에 게이트웨이에서 즉시 거부된다. 실수로 비싼 모델을 반복 호출하는 사고를 원천 차단한다.")

    doc.add_paragraph()

    # 3.3
    heading(doc, "3.3 예산 초과 차단 검증", level=2)
    body(doc,
         "예산 제한이 화면 표시에 그치지 않고 실제 API 호출을 차단하는지 검증하였다. "
         "테스트 키를 생성하여 1회 성공 호출 후 max_budget을 $0.00으로 강제 설정하고 "
         "재호출을 시도하면 즉시 차단된다.")

    body(doc, "실행 명령:", bold=True)
    add_code_block(doc, "uv run python scripts/40_budget_block_test.py")

    body(doc, "검증 흐름:", bold=True)
    for step in [
        "Step 1: 테스트 키 생성 (max_budget=$0.0001)",
        "Step 2: 첫 번째 호출 → 성공 (예산 범위 내)",
        "Step 3: Admin API로 max_budget=$0.00 강제 설정 (예산 소진 시뮬레이션)",
        "Step 4: 두 번째 호출 → 차단 (HTTP 400)",
        "Step 5: 테스트 키 삭제",
    ]:
        bullet(doc, step)

    body(doc, "LiteLLM 응답 (HTTP 400):", bold=True)
    add_code_block(doc,
        '{\n'
        '  "error": {\n'
        '    "message": "Budget has been exceeded! Current cost: 1.35e-05, Max budget: 0.0",\n'
        '    "type": "budget_exceeded",\n'
        '    "code": "400"\n'
        '  }\n'
        '}'
    )

    body(doc,
         "LiteLLM은 호출 단위로 비용을 추적하여 누적 spend가 max_budget에 도달하는 순간 "
         "추가 호출을 즉시 거부한다. 비용은 호출 완료 후 집계되므로 "
         "동시 다발적 호출 환경에서는 짧은 시간 초과가 발생할 수 있다 — 이는 알려진 한계이다.")

    doc.add_paragraph()

    # 3.4
    heading(doc, "3.4 사용량 리포트 검증", level=2)
    body(doc,
         "LiteLLM Admin API의 키 목록 및 spend 로그를 조회하여 "
         "직원별·도구별·모델별 비용을 집계하고 Markdown 리포트로 출력하였다.")

    body(doc, "실행 명령:", bold=True)
    add_code_block(doc, "uv run python scripts/50_export_report.py")

    body(doc, "리포트 주요 항목:", bold=True)
    for item in [
        "가상키 현황 테이블: alias, user, team, tool, max_budget, 실제 spend",
        "사용자별 총 비용: kim / lee / park 분리",
        "도구별 총 비용: claude-code / codex-cli / gemini-cli / chat / admin-api 분리",
        "모델별 호출 수 및 비용: gpt-4o-mini / claude-3-5-sonnet / ssw-fake 등",
        "차단 이벤트 목록: budget_exceeded / key_model_access_denied 이유 포함",
    ]:
        bullet(doc, item)

    doc.add_paragraph()

    # ── 4. Claude Code / Codex CLI 통제 가능성 ─────────────────────────────────
    heading(doc, "4. Claude Code / Codex CLI 가상키 통제 — 검증 상태")

    body(doc,
         "이 질문은 PoC에서 가장 중요한 의문 중 하나이다. 각 CLI 도구가 "
         "LiteLLM Gateway를 경유하도록 설정할 수 있다면, 도구별 가상키로 비용을 완전히 통제할 수 있다.")

    # 4.1
    heading(doc, "4.1 Claude Code — 통제 가능성 (검증 예정)", level=2)
    body(doc, "기술적 원리:", bold=True)
    body(doc,
         "Claude Code는 Anthropic API를 호출할 때 ANTHROPIC_BASE_URL 환경변수로 "
         "엔드포인트를 재지정할 수 있다. LiteLLM Proxy는 Anthropic-compatible API를 "
         "제공하므로 이론적으로 경유 가능하다.")

    add_code_block(doc,
        "# Claude Code가 LiteLLM Gateway를 경유하도록 설정하는 방법 (예시)\n"
        "export ANTHROPIC_BASE_URL=http://localhost:4000\n"
        "export ANTHROPIC_AUTH_TOKEN=sk-JsByY...  # dev-kim-claude 가상키\n"
        "claude  # 이후 Claude Code 실행"
    )

    body(doc, "이번 PoC에서의 검증 범위:", bold=True)
    bullet(doc, "✅ LiteLLM 측 설정 완료: dev-kim-claude 키에 claude-sonnet 모델 권한 및 $5 예산 설정")
    bullet(doc, "✅ 가상키 발급 및 모델 제한 정책 적용 확인")
    bullet(doc, "⚠️ Claude Code 실기 연동 미검증: 실제 Claude Code 바이너리를 통해 ANTHROPIC_BASE_URL을 "
                "LiteLLM으로 지정한 후 호출이 LiteLLM 로그에 기록되는지는 Phase 2에서 검증 예정")
    bullet(doc, "⚠️ Claude Code 버전별 환경변수 명칭 변경 가능성 — 현재 버전 실기 확인 필요")

    body(doc, "성공 기준 (Phase 2):", bold=True)
    body(doc,
         "Claude Code에서 실행한 요청이 LiteLLM usage log에 dev-kim-claude alias로 기록되어야 한다. "
         "이것이 확인되어야 Claude Code가 게이트웨이 통제 하에 있다고 볼 수 있다.")

    doc.add_paragraph()

    # 4.2
    heading(doc, "4.2 Codex CLI — 통제 가능성 (검증 예정)", level=2)
    body(doc, "기술적 원리:", bold=True)
    body(doc,
         "Codex CLI는 OpenAI-compatible custom provider를 지원하며 base_url을 재지정할 수 있다. "
         "LiteLLM은 OpenAI-compatible API를 제공하므로 경유 설정이 가능하다.")

    add_code_block(doc,
        "# Codex CLI가 LiteLLM Gateway를 경유하도록 설정하는 방법 (예시)\n"
        "# ~/.codex/config.yaml 또는 환경변수\n"
        "export OPENAI_BASE_URL=http://localhost:4000\n"
        "export OPENAI_API_KEY=sk-S1nph...  # dev-kim-codex 가상키\n"
        "codex  # 이후 Codex CLI 실행"
    )

    body(doc, "이번 PoC에서의 검증 범위:", bold=True)
    bullet(doc, "✅ LiteLLM 측 설정 완료: dev-kim-codex 키에 gpt-dev 모델 권한 및 $5 예산 설정")
    bullet(doc, "⚠️ Codex CLI 실기 연동 미검증: ChatGPT OAuth 세션이 내장 provider를 우선 사용할 경우 "
                "custom base_url 설정이 무시될 수 있음 — Phase 2에서 검증 필요")
    bullet(doc, "⚠️ Codex CLI의 Responses API 호환성: LiteLLM의 mode: responses 설정 필요 여부 확인 필요")

    doc.add_paragraph()

    # 4.3
    heading(doc, "4.3 Gemini CLI — 통제 가능성 (검증 예정)", level=2)
    body(doc,
         "Gemini CLI도 base URL 환경변수를 통한 재지정이 가능할 것으로 예상되나, "
         "Google 로그인 캐시가 있을 경우 내장 인증 경로를 우선 사용할 가능성이 있다. "
         "dev-kim-gemini 키에 ssw-free-test 모델을 할당해 두었으며, 실기 검증은 Phase 2에서 수행한다.")

    body(doc,
         "\n결론: 세 CLI 도구 모두 LiteLLM Gateway 경유를 위한 가상키와 모델 정책은 이미 준비되어 있다. "
         "Phase 2에서 각 CLI의 실제 바이너리를 통해 호출이 LiteLLM 로그에 기록되는지 확인하면 "
         "통제 가능성을 최종 확정할 수 있다.",
         bold=False)

    doc.add_paragraph()

    # ── 5. PoC 결과의 의미 ────────────────────────────────────────────────────
    heading(doc, "5. PoC 결과의 의미 및 비즈니스 가치")

    heading(doc, "5.1 기술적 의미 — '차단'은 실제로 동작한다", level=2)
    body(doc,
         "이번 PoC에서 가장 중요한 확인 사항은 예산 제한과 모델 제한이 단순한 '설정값 표시'가 아니라 "
         "실제 API 호출을 거부하는 하드 블로킹으로 작동한다는 것이다.")
    for point in [
        "모델 차단은 provider(Anthropic/OpenAI)에 도달하기 전에 게이트웨이 레벨에서 즉시 거부 → 비용 발생 없음",
        "예산 초과 시 LiteLLM이 자동으로 HTTP 400 반환 → 추가 API 키 없이 중앙에서 통제",
        "사용량 로그는 호출 단위로 기록 → 월 결산 시 직원/도구/모델별 비용 정산 가능",
    ]:
        bullet(doc, point)

    heading(doc, "5.2 운영적 의미 — 관리자가 할 수 있는 것", level=2)
    add_table(doc,
        ["관리자 행동", "LiteLLM 기능", "효과"],
        [
            ["직원 입사", "가상키 발급 (Admin API)", "즉시 접근 허용"],
            ["직원 퇴사/이동", "가상키 폐기 (Admin API)", "즉시 접근 차단"],
            ["고가 모델 제한", "키별 allowed_models 설정", "비허용 모델 호출 즉시 차단"],
            ["월 예산 설정", "키별 max_budget 설정", "초과 호출 자동 차단"],
            ["비용 현황 확인", "Admin API 또는 대시보드", "실시간 직원/도구별 비용 조회"],
            ["긴급 차단", "키별 max_budget=0 설정", "즉시 모든 호출 차단"],
        ],
        col_widths=[4, 4.5, 6.5],
    )

    doc.add_paragraph()

    heading(doc, "5.3 전략적 의미 — 왜 이것이 필요한가", level=2)
    body(doc,
         "기업이 직원에게 Anthropic/OpenAI API 키를 직접 발급하면 다음 문제가 발생한다:")
    for p in [
        "키 유출 시 피해 규모를 즉각 파악하기 어렵다",
        "직원별·프로젝트별 비용이 하나의 청구서에 합산되어 ROI 측정 불가",
        "고가 모델 오남용을 기술적으로 막을 방법이 없다",
        "AI 에이전트가 예산 초과 시 자동 중단되지 않는다",
    ]:
        bullet(doc, p)

    body(doc, "\nLiteLLM Gateway 도입 시:")
    for p in [
        "직원은 가상키만 보유 → 실제 provider 키는 중앙에서만 관리",
        "직원이 퇴사해도 가상키만 폐기하면 되고 provider 키 교체 불필요",
        "도구별·직원별·프로젝트별 비용을 별도 집계하여 ROI 산정 가능",
        "예산 초과 자동 차단으로 AI 비용 서프라이즈 방지",
        "향후 Google Chat, Hermes Agent 등 모든 도구를 단일 게이트웨이로 통제",
    ]:
        bullet(doc, p)

    doc.add_paragraph()

    # ── 6. 한계 및 주의사항 ────────────────────────────────────────────────────
    heading(doc, "6. 한계 및 주의사항")

    add_table(doc,
        ["항목", "현황", "대응 방안"],
        [
            ["개인 ChatGPT/Claude 앱 통제",
             "기술적으로 불가 (개인 계정 기반)",
             "정책으로 금지 + 업무용 게이트웨이 의무 사용 규정"],
            ["Claude Code / Codex / Gemini CLI 연동",
             "PoC에서 미검증 (Phase 2)",
             "각 CLI의 base URL 환경변수 실기 확인 필요"],
            ["RPM/TPM 속도 제한",
             "Redis 미설정으로 비활성",
             "Redis 추가 시 즉시 활성화 가능"],
            ["예산 차단 타이밍",
             "호출 완료 후 집계 → 동시 호출 시 짧은 초과 가능",
             "운영 시 허용 오차 범위를 예산에 반영"],
            ["운영 가용성",
             "LiteLLM 서버 장애 시 모든 AI 호출 중단",
             "HA 구성, 헬스체크, 백업 필요"],
            ["SSO 연동",
             "PoC에서 미포함",
             "Phase 3에서 Google Workspace SSO 연동"],
        ],
        col_widths=[4, 5, 6],
    )

    doc.add_paragraph()

    # ── 7. 다음 단계 ───────────────────────────────────────────────────────────
    heading(doc, "7. 권장 다음 단계")

    add_table(doc,
        ["단계", "작업", "우선순위"],
        [
            ["Phase 2-A", "Claude Code ANTHROPIC_BASE_URL 실기 연동 검증", "즉시"],
            ["Phase 2-B", "Codex CLI custom provider 실기 연동 검증", "즉시"],
            ["Phase 2-C", "OpenRouter 무료 API 키로 실제 LLM 호출 검증", "즉시"],
            ["Phase 3", "Redis 설치 → RPM/TPM 속도 제한 활성화", "단기"],
            ["Phase 4", "Hermes Agent LiteLLM provider 연동", "단기"],
            ["Phase 5", "Google Chat 업무봇 연동", "중기"],
            ["Phase 6", "SSO 연동 + 키 로테이션 + 감사 로그", "중기"],
            ["Phase 7", "ROI 엔진 + 관리자 대시보드 완성", "장기"],
        ],
        col_widths=[2.5, 9, 2.5],
    )

    doc.add_paragraph()

    # ── 8. 결론 ───────────────────────────────────────────────────────────────
    heading(doc, "8. 결론")

    body(doc,
         "본 PoC는 LiteLLM Proxy를 AI 비용 게이트웨이로 활용하여 "
         "직원별·도구별 LLM 사용을 기업 예산 범위 안에서 통제할 수 있음을 기술적으로 증명하였다.",
         bold=True)

    doc.add_paragraph()

    body(doc, "핵심 결론 3가지:")
    for c in [
        "✅ 가상키별 모델 제한은 실제로 작동한다 — 비허용 모델 호출은 provider 도달 전에 게이트웨이에서 거부된다.",
        "✅ 예산 초과 차단은 실제로 작동한다 — max_budget 도달 즉시 API 호출이 HTTP 400으로 거부된다.",
        "✅ 사용량 추적은 가능하다 — 직원/도구/모델 단위로 비용이 분리 기록되어 리포트 생성이 가능하다.",
    ]:
        bullet(doc, c)

    doc.add_paragraph()

    body(doc,
         "Claude Code, Codex CLI, Gemini CLI의 실기 연동은 Phase 2에서 검증 예정이며, "
         "기술적 구조상 각 도구의 base URL 환경변수를 LiteLLM으로 지정하면 동일한 "
         "모델 제한 및 예산 차단 정책이 적용될 것으로 예상된다. "
         "이 검증이 완료되면 상상우리 AI 에이전트 프로젝트의 비용 거버넌스 기반이 확립된다.")

    # ── 저장 ──────────────────────────────────────────────────────────────────
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"✅ 보고서 저장 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
